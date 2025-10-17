# app.py — Crushers parts by Asset (no Excel formulas in sheet required)

import io, os, re
from pathlib import Path
from datetime import datetime
from zoneinfo import ZoneInfo

import pandas as pd
import streamlit as st

# ==================== Config ====================
DATA_PATH   = Path("Crushers.xlsx")   # workbook in repo root
SHEET_NAME  = None                    # or set a specific sheet name
MANUALS_DIR = Path("manuals")         # put <Asset>.pdf here

# Optional GitHub raw link (set to enable a clickable link)
GH_OWNER  = ""       # e.g., "youruser"
GH_REPO   = ""       # e.g., "crushers-repo"
GH_BRANCH = "main"

st.set_page_config(page_title="Crushers Parts", layout="wide")
st.title("Crushers Parts")

# Try to import python-docx for Word export
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ==================== Helpers ====================
def pick_col(df, candidates):
    """Return first matching column (case/spacing tolerant)."""
    norm = {c: re.sub(r"\W+", "", str(c).lower()) for c in df.columns}
    cands = [re.sub(r"\W+", "", x.lower()) for x in candidates]
    for c, n in norm.items():
        if n in cands:
            return c
    for cand in cands:
        for c in df.columns:
            if cand in str(c).lower():
                return c
    return None

def to_int_or_nan(v):
    try:
        return int(float(v))
    except Exception:
        return pd.NA

def sanitize_filename(name: str) -> str:
    return re.sub(r"[^\w\-]+", "_", str(name)).strip("_")

def normalize_fname(s: str) -> str:
    # lower + unify separators for tolerant matching
    return re.sub(r"[^\w\-]+", "_", str(s)).strip("_").lower().replace("-", "_")

def find_manual_path(asset_name: str) -> Path | None:
    """Find <Asset>.pdf in manuals/ (case/spacing tolerant)."""
    if not MANUALS_DIR.exists():
        return None
    target = normalize_fname(asset_name)
    for p in MANUALS_DIR.glob("*.pdf"):
        if normalize_fname(p.stem) == target:
            return p
    return None

def github_raw_url(asset_name: str) -> str:
    if not (GH_OWNER and GH_REPO):
        return ""
    rel = (MANUALS_DIR / (sanitize_filename(asset_name) + ".pdf")).as_posix()
    return f"https://raw.githubusercontent.com/{GH_OWNER}/{GH_REPO}/{GH_BRANCH}/{rel}"

def to_excel_bytes(dfout: pd.DataFrame, sheet_name="Parts") -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="xlsxwriter") as xw:
        dfout.to_excel(xw, index=False, sheet_name=sheet_name)
        ws = xw.sheets.get(sheet_name)
        if ws:
            # widths tuned for: Page | Item No | QTY | Part No | Part Name | InStk
            for idx, width in enumerate([10, 14, 8, 18, 36, 30], start=1):
                try:
                    ws.set_column(idx-1, idx-1, width)
                except Exception:
                    pass
    bio.seek(0)
    return bio.getvalue()

# ========= DOCX Export (updated to Mfr/Model/Serial + new columns) =========
def word_bytes_from_rows(dfq: pd.DataFrame, df_all: pd.DataFrame,
                         col_asset: str, col_model: str | None, col_serial: str | None) -> bytes:
    """
    Create a .docx with:
      - Heading + Generated timestamp
      - Vendor placeholder line
      - Manufacturer/Model/Serial line(s) under Vendor (one per asset; skips blanks)
      - Parts table: Page • Item • Part Number • Part Name • Qty (+ 10 blank rows)
    """
    if not DOCX_AVAILABLE or dfq.empty:
        return b""

    # Figure out Manufacturer column name
    mfr_col = pick_col(df_all, ["Manufacturer", "Mfr", "MFG", "Make", "Brand"])

    # Collect unique, non-null assets from selected rows
    assets = sorted(dfq["Asset"].dropna().astype(str).unique().tolist())

    # Build per-asset meta rows; dedupe identical lines
    meta_set = set()
    meta_rows = []
    for a in assets:
        mfr = model = serial = ""
        try:
            rows = df_all[df_all[col_asset].astype(str) == a]
            if not rows.empty:
                if mfr_col and pd.notna(rows.iloc[0][mfr_col]):    mfr    = str(rows.iloc[0][mfr_col])
                if col_model and pd.notna(rows.iloc[0][col_model]):  model  = str(rows.iloc[0][col_model])
                if col_serial and pd.notna(rows.iloc[0][col_serial]): serial = str(rows.iloc[0][col_serial])
        except Exception:
            pass
        trip = (mfr.strip(), model.strip(), serial.strip())
        if any(trip) and trip not in meta_set:
            meta_set.add(trip)
            meta_rows.append({"Manufacturer": trip[0], "Model": trip[1], "Serial": trip[2]})

    # Build parts table with requested columns
    def col_or_blank(df: pd.DataFrame, name: str) -> pd.Series:
        return df[name] if name in df.columns else pd.Series([""] * len(df), index=df.index)

    parts = pd.DataFrame(index=dfq.index)
    parts["Page"]         = col_or_blank(dfq, "Page").astype(str)
    parts["Item"]         = col_or_blank(dfq, "Item Number").astype(str)
    parts["Part Number"]  = col_or_blank(dfq, "Part Number").astype(str)
    parts["Part Name"]    = col_or_blank(dfq, "Part Name").astype(str)
    parts["Qty"]          = col_or_blank(dfq, "QTY").astype(str)

    # Add 10 blank rows for manual add-ons
    blanks = pd.DataFrame([{"Page": "", "Item": "", "Part Number": "", "Part Name": "", "Qty": ""} for _ in range(10)])
    parts  = pd.concat([parts.reset_index(drop=True), blanks], ignore_index=True)

    # Build doc
    doc = Document()
    try:
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)
    except Exception:
        pass

    doc.add_heading("Quote Request", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Vendor placeholder
    p = doc.add_paragraph("Vendor: ")
    r = p.add_run("Enter vendor here")
    r.italic = True

    # Manufacturer / Model / Serial lines (under vendor)
    if meta_rows:
        doc.add_paragraph("")  # spacing
        for row in meta_rows:
            m = row.get("Manufacturer", "").strip()
            mo = row.get("Model", "").strip()
            se = row.get("Serial", "").strip()
            doc.add_paragraph(f"Manufacturer: {m}    Model: {mo}    Serial: {se}")

    doc.add_paragraph("")  # spacing
    doc.add_heading("Requested Parts", level=2)

    # Table with the 5 requested columns
    cols = ["Page", "Item", "Part Number", "Part Name", "Qty"]
    t = doc.add_table(rows=1, cols=len(cols))
    hdr = t.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
    for _, rr in parts.iterrows():
        cells = t.add_row().cells
        cells[0].text = rr["Page"]
        cells[1].text = rr["Item"]
        cells[2].text = rr["Part Number"]
        cells[3].text = rr["Part Name"]
        cells[4].text = rr["Qty"]

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()

@st.cache_data
def load_data(_data_path: str, _sheet_name: str | None):
    if _sheet_name:
        df = pd.read_excel(_data_path, sheet_name=_sheet_name)
        used = _sheet_name
    else:
        xls = pd.ExcelFile(_data_path)
        df, used = None, None
        for sh in xls.sheet_names:
            t = pd.read_excel(xls, sheet_name=sh)
            if any("instk" in str(c).lower() or "in stock" in str(c).lower() for c in t.columns):
                df, used = t, sh
                break
        if df is None:
            df = pd.read_excel(xls, sheet_name=0)
            used = xls.sheet_names[0]
    return df, used

def file_last_updated_label(path: Path) -> str:
    try:
        ts = path.stat().st_mtime
        dt_et = datetime.fromtimestamp(ts, tz=ZoneInfo("America/New_York"))
        return dt_et.strftime("%Y-%m-%d %H:%M ") + dt_et.tzname()  # EST/EDT
    except Exception:
        return "unknown"

# ==================== Load sheet ====================
try:
    df, used_sheet = load_data(str(DATA_PATH), SHEET_NAME)
    last = file_last_updated_label(DATA_PATH)
    st.caption(f"Data file: **{DATA_PATH.name}** • last updated: **{last}** • sheet: **{used_sheet}**")
except Exception as e:
    st.error(f"Could not open `{DATA_PATH}`: {e}")
    st.stop()

# ==================== Column mapping ====================
col_asset = pick_col(df, ["Asset", "Asset #", "Asset ID", "Equipment", "Equipment #", "Equipment ID", "Machine"])
col_page  = pick_col(df, ["Page", "Pg"])
col_item  = pick_col(df, ["Item Number", "Item #", "Item No", "Item", "Ref", "Ref #"])
col_qty   = pick_col(df, ["QTY", "Qty", "Quantity"])
col_pn    = pick_col(df, ["Part Number", "Part Numbers", "PN"])
col_name  = pick_col(df, ["Part Name", "Name", "Description", "Part Description"])
col_instk = pick_col(df, ["InStk", "In Stock", "LocAreaQty", "Loc Area Qty", "Location: Area,; Qty", "In_Stock"])
col_model = pick_col(df, ["Model", "Machine Model"])
col_serial= pick_col(df, ["Serial", "Serial Number", "S/N", "SN"])

needed = {
    "Asset": col_asset, "Page": col_page, "Item Number": col_item,
    "QTY": col_qty, "Part Number": col_pn, "Part Name": col_name, "InStk": col_instk
}
missing = [n for n, c in needed.items() if c is None]
if missing:
    st.error(f"Missing required column(s): {', '.join(missing)}\n\nColumns found: {list(df.columns)}")
    st.stop()

# ==================== Top controls (compact) ====================
assets = sorted(df[col_asset].dropna().astype(str).unique().tolist())

# Row: Asset | Page | Page filter | Search parts
c_asset, c_page, c_pagefilter, c_search = st.columns([1.6, 1.0, 0.9, 2.0])
with c_asset:
    sel_asset = st.selectbox("Asset", assets, index=0)

pages_for_asset = df.loc[df[col_asset].astype(str)==str(sel_asset), col_page].dropna().astype(str).unique().tolist()
pages_for_asset = sorted(pages_for_asset, key=lambda x: (to_int_or_nan(x), x))

with c_pagefilter:
    page_filter = st.text_input("Page filter", value="", placeholder="e.g. 42")

page_opts = [p for p in pages_for_asset if page_filter.lower() in p.lower()] if page_filter else pages_for_asset
with c_page:
    sel_page = st.selectbox("Page", ["All"] + page_opts, index=0)

with c_search:
    search_parts = st.text_input("Search parts", value="", placeholder="Part number or name")

# Manual links/downloads
manual_local  = find_manual_path(sel_asset)
manual_expect = (MANUALS_DIR / (sanitize_filename(sel_asset) + ".pdf")).as_posix()
cols_top = st.columns([1, 1, 6])
with cols_top[0]:
    if manual_local and manual_local.exists():
        try:
            pdf_bytes = manual_local.read_bytes()
            st.download_button("Manual PDF", data=pdf_bytes,
                               file_name=manual_local.name, mime="application/pdf")
        except Exception as e:
            st.warning(f"Manual found but couldn't read: {e}")
    else:
        st.caption(f"No local manual. Expected: `{manual_expect}`")
with cols_top[1]:
    gh_url = github_raw_url(sel_asset)
    if gh_url:
        try:
            st.link_button("Manual (GitHub raw)", gh_url)
        except Exception:
            st.markdown(f"[Manual (GitHub raw)]({gh_url})")

# ==================== View ====================
f = (df[col_asset].astype(str) == str(sel_asset))
if sel_page != "All":
    f &= (df[col_page].astype(str) == str(sel_page))

# Keep Asset internally, but don't display it
view = df.loc[f, [col_asset, col_page, col_item, col_qty, col_pn, col_name, col_instk]].copy()
view.rename(columns={
    col_asset: "Asset",
    col_page: "Page",
    col_item: "Item Number",
    col_qty: "QTY",
    col_pn: "Part Number",
    col_name: "Part Name",
    col_instk: "InStk"
}, inplace=True)

if search_parts:
    sp = search_parts.strip().lower()
    view = view[
        view["Part Number"].astype(str).str.lower().str.contains(sp, na=False) |
        view["Part Name"].astype(str).str.lower().str.contains(sp, na=False)
    ]

# Sort (numeric-friendly)
view["_page_sort"] = view["Page"].map(to_int_or_nan)
view["_item_sort"] = view["Item Number"].map(to_int_or_nan)
view = view.sort_values(by=["_page_sort", "Page", "_item_sort", "Item Number"]).drop(columns=["_page_sort", "_item_sort"])

# ---------- Display table WITHOUT Asset ----------
view_disp = view.drop(columns=["Asset"]).copy()
view_disp.insert(0, "Select", False)

st.subheader(f"Parts for: {sel_asset}")
edited = st.data_editor(
    view_disp,
    hide_index=True,
    use_container_width=True,
    column_config={
        "Select": st.column_config.CheckboxColumn("Select"),
        "Page": st.column_config.TextColumn("Page"),
        "Item Number": st.column_config.TextColumn("Item Number"),
        "QTY": st.column_config.NumberColumn("QTY"),
        "Part Number": st.column_config.TextColumn("Part Number"),
        "Part Name": st.column_config.TextColumn("Part Name"),
        "InStk": st.column_config.TextColumn("InStk"),
    },
    disabled=["Page", "Item Number", "QTY", "Part Number", "Part Name", "InStk"]
)

# ==================== Quote cart (multi-page / multi-asset) ====================
# edited doesn't include Asset; map selected rows back to 'view' to reattach Asset
sel_rows_now_display = edited[edited["Select"]].drop(columns=["Select"]).copy()

# Merge back to recover Asset (use a robust key set)
merge_keys = ["Page","Item Number","Part Number","Part Name","QTY","InStk"]
sel_rows_now = pd.merge(sel_rows_now_display, view, on=merge_keys, how="left")

if "quote_cart" not in st.session_state:
    st.session_state.quote_cart = pd.DataFrame(columns=["Asset","Page","Item Number","Part Number","Part Name","QTY","InStk"])

def add_to_cart(df_add: pd.DataFrame):
    if df_add.empty:
        return
    cart = st.session_state.quote_cart
    combined = pd.concat([cart, df_add], ignore_index=True)
    # De-dupe by Asset + Page + Item Number + Part Number
    dedupe_keys = ["Asset","Page","Item Number","Part Number"]
    combined = combined.sort_values(dedupe_keys).drop_duplicates(subset=dedupe_keys, keep="last")
    st.session_state.quote_cart = combined

c_cart1, c_cart2, c_cart3, _ = st.columns([1.2, 1.6, 1.2, 5])
with c_cart1:
    if st.button("Add selected to cart", type="primary"):
        add_to_cart(sel_rows_now)
with c_cart2:
    if st.button("Clear cart"):
        st.session_state.quote_cart = st.session_state.quote_cart.iloc[0:0].copy()
with c_cart3:
    st.caption(f"Cart items: **{len(st.session_state.quote_cart)}**")

# ---- Editable Cart (change QTY, remove items) ----
st.markdown("### Cart")
cart_df = st.session_state.quote_cart.copy()

if cart_df.empty:
    st.caption("No items yet. Select rows above and click **Add selected to cart**.")
else:
    # Build a user-facing cart view; keep Asset internally but show for context
    cart_view = cart_df.copy()
    cart_view.insert(0, "Remove", False)

    edited_cart = st.data_editor(
        cart_view[["Remove","Asset","Page","Item Number","Part Number","Part Name","QTY","InStk"]],
        hide_index=True,
        use_container_width=True,
        column_config={
            "Remove": st.column_config.CheckboxColumn("Remove", help="Check to remove selected rows"),
            "QTY": st.column_config.NumberColumn("QTY", step=1, min_value=0),
            "Asset": st.column_config.TextColumn("Asset"),
            "Page": st.column_config.TextColumn("Page"),
            "Item Number": st.column_config.TextColumn("Item Number"),
            "Part Number": st.column_config.TextColumn("Part Number"),
            "Part Name": st.column_config.TextColumn("Part Name"),
            "InStk": st.column_config.TextColumn("InStk"),
        },
        disabled=["Asset","Page","Item Number","Part Number","Part Name","InStk"],
        key="cart_editor",
    )

    # Buttons on one line: Remove (left) • [spacer] • Clear • Save • Generate (right)
    c_remove, c_spacer, c_clear, c_save, c_gen = st.columns([1.1, 6.5, 0.9, 0.9, 1.2])
    remove_it = c_remove.button("Remove", use_container_width=True)
    clear_it  = c_clear.button("Clear", use_container_width=True)
    save_qty  = c_save.button("Save", use_container_width=True)

    # Generate uses the *current edited view* so Save isn't strictly required
    rows_for_quote = edited_cart.drop(columns=["Remove"], errors="ignore").copy()
    if not rows_for_quote.empty and DOCX_AVAILABLE:
        quote_blob = word_bytes_from_rows(rows_for_quote, df, col_asset, col_model, col_serial)
        base_name  = f"QuoteRequest_{sanitize_filename(str(sel_asset))}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        c_gen.download_button(
            "Generate",
            data=quote_blob,
            file_name=base_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        c_gen.download_button(
            "Generate",
            data=b"",
            file_name="QuoteRequest.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            disabled=True
        )

    # Apply actions after rendering buttons
    if save_qty and "QTY" in edited_cart.columns:
        st.session_state.quote_cart.loc[edited_cart.index, "QTY"] = edited_cart["QTY"].values
        st.success("Saved quantities.")

    if remove_it:
        try:
            idx = edited_cart.index[edited_cart["Remove"] == True]
        except Exception:
            idx = []
        if len(idx):
            st.session_state.quote_cart = st.session_state.quote_cart.drop(index=idx).reset_index(drop=True)
            st.rerun()

    if clear_it:
        st.session_state.quote_cart = st.session_state.quote_cart.iloc[0:0].copy()
        st.rerun()

with st.expander("View raw cart data"):
    st.dataframe(st.session_state.quote_cart, hide_index=True, use_container_width=True)

# ==================== Downloads ====================
# Download Current View (Excel) — exports what you see (NO Asset column)
current_view_xlsx = to_excel_bytes(view_disp.drop(columns=["Select"]), sheet_name="Parts")
st.download_button("Download Current View (Excel)",
                   data=current_view_xlsx,
                   file_name=f"Parts_{sanitize_filename(str(sel_asset))}.xlsx",
                   mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")




